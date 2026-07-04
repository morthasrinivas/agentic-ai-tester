"""
Document loader: reads PDF and DOCX files and returns plain text chunks.
"""

from __future__ import annotations
from pathlib import Path
from typing import List
import re


def _clean(text: str) -> str:
    """Normalise whitespace and remove page artefacts."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def load_docx(path: Path) -> str:
    """Extract full text from a .docx file."""
    from docx import Document
    doc = Document(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return _clean("\n".join(parts))


def load_pdf(path: Path) -> str:
    """Extract full text from a .pdf file using pdfplumber."""
    import pdfplumber
    parts: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return _clean("\n".join(parts))


def load_document(path: Path) -> str:
    """Load either a PDF or DOCX file and return raw text."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx(path)
    elif suffix == ".pdf":
        return load_pdf(path)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    """
    Split text into overlapping chunks for embedding.
    Tries to split on paragraph boundaries first.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= chunk_size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            # If a single paragraph exceeds chunk_size, split by character
            if len(para) > chunk_size:
                for i in range(0, len(para), chunk_size - overlap):
                    chunks.append(para[i: i + chunk_size])
            else:
                current = para

    if current:
        chunks.append(current)

    return chunks
